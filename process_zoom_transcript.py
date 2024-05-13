#!/usr/bin/env python
import os
import argparse
import re

import docx

def parseArguments():
    parser = argparse.ArgumentParser()
    parser.add_argument("file", help="Transcript file", type=str)
    parser.add_argument("-s", "--speaker", help="Speaker to retain", type=str, action="append")
    parser.add_argument("output", help="Output file name", type=str)

    args = parser.parse_args()
    return args

def parse_text(args):
    doc = docx.Document()

    transcript = open(args.file, 'r')

    current_speaker = ""
    count = 0
    current_paragraph = ""

    for line in transcript:
        count += 1
        for speaker in args.speaker:
            if line.startswith(speaker):
                if speaker == current_speaker:
                    current_paragraph.add_run(" " + re.sub(r'^.*?:', '', line).strip())
                else:
                    current_speaker = speaker
                    current_paragraph = doc.add_paragraph(line.strip())

    doc.save(args.output)

def main(args):
    parse_text(args)
    print('Output file:', args.output)

if __name__ == '__main__':
    args = parseArguments()
    main(args)
